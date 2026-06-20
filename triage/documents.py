from __future__ import annotations

import base64
import logging
from pathlib import Path

from docx import Document
from fastapi import UploadFile

from triage.models import UploadedDocument

logger = logging.getLogger("triage.documents")


PDF_TYPES = {"application/pdf"}
DOCX_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/octet-stream",
}
TXT_TYPES = {"text/plain", "text/markdown"}


def _guess_mime(filename: str, declared: str | None) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return "application/pdf"
    if suffix == ".docx":
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if suffix in {".txt", ".md"}:
        return "text/plain"
    return declared or "application/octet-stream"


def _docx_text(raw: bytes) -> str:
    import io

    logger.info("DOC_PARSE docx extraction started bytes=%s", len(raw))
    doc = Document(io.BytesIO(raw))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    logger.info("DOC_PARSE docx extraction complete paragraphs=%s chars=%s", len(paragraphs), len("\n".join(paragraphs)))
    return "\n".join(paragraphs)


async def read_uploads(files: list[UploadFile]) -> list[UploadedDocument]:
    docs: list[UploadedDocument] = []
    logger.info("DOC_READ started file_count=%s", len(files))
    for f in files:
        filename = f.filename or "document"
        logger.info("DOC_READ file start filename=%s declared_type=%s", filename, f.content_type or "-")
        raw = await f.read()
        mime = _guess_mime(filename, f.content_type)
        text: str | None = None
        if mime in TXT_TYPES:
            text = raw.decode("utf-8", errors="replace")
            logger.info("DOC_READ text decoded filename=%s mime=%s chars=%s", filename, mime, len(text))
        elif mime in DOCX_TYPES and (f.filename or "").lower().endswith(".docx"):
            text = _docx_text(raw)
            logger.info("DOC_READ docx decoded filename=%s mime=%s chars=%s", filename, mime, len(text))
        elif mime not in PDF_TYPES:
            text = raw.decode("utf-8", errors="replace")
            logger.info("DOC_READ binary-as-text decoded filename=%s mime=%s chars=%s", filename, mime, len(text))
        else:
            logger.info("DOC_READ pdf captured filename=%s mime=%s bytes=%s", filename, mime, len(raw))
        docs.append(
            UploadedDocument(
                filename=filename,
                mime_type=mime,
                size=len(raw),
                text=text,
                raw_bytes_b64=base64.b64encode(raw).decode("ascii"),
            )
        )
        logger.info(
            "DOC_READ file complete filename=%s mime=%s bytes=%s has_text=%s",
            filename,
            mime,
            len(raw),
            bool(text),
        )
    logger.info("DOC_READ complete parsed_count=%s", len(docs))
    return docs


def safe_document_dump(docs: list[dict]) -> list[dict]:
    safe = []
    for d in docs:
        item = dict(d)
        item.pop("raw_bytes_b64", None)
        safe.append(item)
    return safe
